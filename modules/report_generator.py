from docx import Document
from docx.shared import Pt
from datetime import datetime
import os
import subprocess


# Funciones para generar los comentarios (como se mostró antes)
def generate_comment_IMC(patient):
    if patient.imc >= 30:
        return f"El IMC del paciente es de {patient.imc}, lo cual indica obesidad. Es recomendable un plan de reducción de peso."
    elif patient.imc >= 25:
        return f"El IMC del paciente es de {patient.imc}, lo cual indica sobrepeso. Se recomienda mejorar la dieta y aumentar la actividad física."
    else:
        return f"El IMC del paciente es de {patient.imc}, lo cual está dentro del rango normal."

def generate_comment_Presion(patient):
    if patient.presion_sistolica >= 140 or patient.presion_diastolica >= 90:
        return f"La presión arterial del paciente es {patient.presion_sistolica}/{patient.presion_diastolica}, lo cual indica hipertensión. Se recomienda tratamiento y seguimiento."
    else:
        return f"La presión arterial del paciente es {patient.presion_sistolica}/{patient.presion_diastolica}, dentro del rango normal."

def generate_comment_Colesterol(patient):
    if patient.colesterol_total >= 240:
        return f"El colesterol total del paciente es {patient.colesterol_total} mg/dL, lo cual indica hipercolesterolemia. Se recomienda ajustar la dieta y considerar medicación."
    elif patient.colesterol_total >= 200:
        return f"El colesterol total del paciente es {patient.colesterol_total} mg/dL, lo cual está en el límite superior. Se recomienda monitoreo y control dietético."
    else:
        return f"El colesterol total del paciente es {patient.colesterol_total} mg/dL, lo cual está dentro del rango normal."

def generate_comment_Trigliceridos(patient):
    if patient.trigliceridos > 150:
        return f"El nivel de triglicéridos del paciente es {patient.trigliceridos} mg/dL, lo cual está elevado. Se recomienda un control en la dieta y aumentar la actividad física."
    else:
        return f"El nivel de triglicéridos del paciente es {patient.trigliceridos} mg/dL, dentro del rango normal."

def generate_comment_IndiceAterogenico(patient):
    if patient.ia >= 4.5:
        return f"El índice aterogénico del paciente es de {patient.ia}, lo cual indica un riesgo elevado de enfermedades cardiovasculares. Se recomienda un monitoreo más cercano y tratamiento preventivo."
    else:
        return f"El índice aterogénico del paciente es de {patient.ia}, dentro de un rango saludable."

def generate_comment_Glucosa(patient):
    if patient.glucosa >= 126:
        return f"El nivel de glucosa en sangre del paciente es {patient.glucosa} mg/dL, lo cual indica un posible riesgo de diabetes. Se recomienda realizar un seguimiento continuo."
    else:
        return f"El nivel de glucosa en sangre del paciente es {patient.glucosa} mg/dL, dentro del rango normal."

def generate_comment_HbA1c(patient):
    if patient.HbA1c >= 6.5:
        return f"El nivel de HbA1c del paciente es de {patient.HbA1c}%, lo cual está en el rango para diabetes tipo 2. Se recomienda un control adecuado de la glucosa."
    else:
        return f"El nivel de HbA1c del paciente es de {patient.HbA1c}%, dentro del rango saludable."

def generate_comment_Fumador(patient):
    if patient.fumador:
        return f"El paciente es fumador, lo que incrementa el riesgo de enfermedades respiratorias y cardiovasculares. Se recomienda cesación del tabaquismo."
    else:
        return f"El paciente no es fumador, lo cual es positivo para su salud general."

def generate_comment_Alcoholico(patient):
    if patient.alcoholico:
        return f"El paciente tiene antecedentes de alcoholismo, lo cual puede afectar el funcionamiento hepático y cardiovascular. Se recomienda tratamiento y abstención."
    else:
        return f"El paciente no es alcohólico, lo cual es positivo para la salud general."

def generate_comment_Activo(patient):
    if patient.activo:
        return f"El paciente mantiene un estilo de vida activo, lo cual es beneficioso para la salud cardiovascular y metabólica."
    else:
        return f"El paciente no lleva un estilo de vida suficientemente activo. Se recomienda aumentar la actividad física para mejorar la salud general."

def generate_comment_FamiliaresDiabetes(patient):
    if patient.familiares_diabetes:
        return f"El paciente tiene antecedentes familiares de diabetes, lo que aumenta su riesgo de desarrollar la enfermedad. Se recomienda monitoreo constante."
    else:
        return f"El paciente no tiene antecedentes familiares de diabetes, lo cual es un factor positivo."

def generate_comment_FamiliaresHipertension(patient):
    if patient.familiares_hipertension:
        return f"El paciente tiene antecedentes familiares de hipertensión, lo que incrementa su riesgo. Se recomienda monitoreo regular de la presión arterial."
    else:
        return f"El paciente no tiene antecedentes familiares de hipertensión, lo cual es un factor positivo."

def generate_comment_FamiliaresCardiovascular(patient):
    if patient.familiares_cardiovascular:
        return f"El paciente tiene antecedentes familiares de enfermedades cardiovasculares, lo que aumenta su riesgo. Se recomienda realizar un control cardiovascular."
    else:
        return f"El paciente no tiene antecedentes familiares de enfermedades cardiovasculares, lo cual es positivo."

def generate_comment_FamiliaresCerebrovascular(patient):
    if patient.familiares_cerebrovascular:
        return f"El paciente tiene antecedentes familiares de enfermedades cerebrovasculares, lo cual aumenta su riesgo. Se recomienda monitoreo adecuado."
    else:
        return f"El paciente no tiene antecedentes familiares de enfermedades cerebrovasculares, lo cual es positivo."


# Función para generar conclusiones
def generate_conclusions(patient):
    conclusions = []
    
    # Riesgo de diabetes
    if patient.glucosa >= 126 or patient.HbA1c >= 6.5:
        conclusions.append("El paciente presenta un alto riesgo de desarrollar diabetes tipo 2 debido a los niveles elevados de glucosa y HbA1c.")
    
    # Riesgo cardiovascular
    if patient.imc >= 30 and patient.presion_sistolica > 140 and patient.colesterol_total > 240:
        conclusions.append("El paciente tiene un alto riesgo cardiovascular debido a la obesidad, hipertensión y niveles elevados de colesterol.")
    elif patient.imc >= 30 and patient.presion_sistolica > 140:
        conclusions.append("El paciente tiene un alto riesgo cardiovascular debido a la obesidad e hipertensión.")
    elif patient.imc >= 30 and patient.colesterol_total > 240:
        conclusions.append("El paciente tiene un alto riesgo cardiovascular debido a la obesidad y niveles elevados de colesterol.")
    elif patient.presion_sistolica > 140 and patient.colesterol_total > 240:
        conclusions.append("El paciente tiene un alto riesgo cardiovascular debido a la hipertensión y niveles elevados de colesterol.")
    elif patient.imc >= 30:
        conclusions.append("El paciente tiene un alto riesgo cardiovascular debido a la obesidad.")
    elif patient.colesterol_total > 240:
        conclusions.append("El paciente tiene un alto riesgo cardiovascular debido a los niveles elevados de colesterol.")
    elif patient.presion_sistolica > 140:
        conclusions.append("El paciente tiene un alto riesgo cardiovascular debido a la hipertensión.")

    # Riesgo de hipertensión
    if patient.presion_sistolica > 140 or patient.presion_diastolica > 90:
        conclusions.append("El paciente presenta hipertensión, lo que aumenta el riesgo de enfermedades cardiovasculares.")
    
    # Riesgo de enfermedades crónicas por hábitos
    if patient.fumador:
        conclusions.append("El paciente es fumador, lo cual incrementa su riesgo de enfermedades respiratorias y cardiovasculares.")
    
    if not patient.activo:
        conclusions.append("El paciente lleva un estilo de vida sedentario, lo que aumenta el riesgo de enfermedades metabólicas y cardiovasculares.")
    
    if patient.alcoholico:
        conclusions.append("El paciente tiene antecedentes de alcoholismo, lo cual afecta negativamente la salud hepática y cardiovascular.")
    
    # Riesgo general debido a antecedentes familiares
    if patient.familiares_diabetes:
        conclusions.append("El paciente tiene antecedentes familiares de diabetes, lo que incrementa su riesgo de desarrollar esta enfermedad.")
    
    if patient.familiares_hipertension:
        conclusions.append("El paciente tiene antecedentes familiares de hipertensión, lo que aumenta su riesgo de padecerla.")
    
    return "\n\n".join(conclusions)

# Función para generar recomendaciones
def generate_recommendations(patient):
    recommendations = []
    
    # Control de diabetes
    if patient.glucosa >= 126:
        recommendations.append("Se recomienda monitorear los niveles de glucosa y seguir un plan de tratamiento para prevenir la diabetes tipo 2.")
    if patient.HbA1c >= 6.5:
        recommendations.append("Es recomendable un control adecuado de la glucosa y un seguimiento médico regular para evitar complicaciones.")
    
    # Control cardiovascular
    if patient.imc >= 30:
        recommendations.append("Se recomienda seguir un plan de pérdida de peso a través de una dieta balanceada y ejercicio regular.")
    if patient.presion_sistolica > 140 or patient.presion_diastolica > 90:
        recommendations.append("Se recomienda controlar la hipertensión mediante medicación y modificaciones en el estilo de vida, como reducir el estrés y hacer ejercicio.")
    if patient.colesterol_total > 240:
        recommendations.append("Se recomienda cambiar a una dieta baja en grasas saturadas y colesterol, y seguir un plan de ejercicio regular.")
    
    # Hábitos saludables
    if patient.fumador:
        recommendations.append("Es esencial dejar de fumar para reducir el riesgo de enfermedades respiratorias y cardiovasculares. Se recomienda consultar con un especialista para cesación del tabaquismo.")
    
    if not patient.activo:
        recommendations.append("Se recomienda aumentar la actividad física para mejorar la salud cardiovascular y metabólica. Se sugiere realizar al menos 30 minutos de ejercicio moderado al día.")
    
    if patient.alcoholico:
        recommendations.append("Se recomienda abstenerse del consumo de alcohol, o buscar ayuda para superar el alcoholismo, lo cual mejorará la salud hepática y cardiovascular.")
    
    # Control de antecedentes familiares
    if patient.familiares_diabetes:
        recommendations.append("Dado el riesgo hereditario, es fundamental un monitoreo regular de los niveles de glucosa y un estilo de vida saludable.")
    
    if patient.familiares_hipertension:
        recommendations.append("Se recomienda monitorear la presión arterial con regularidad y adoptar hábitos saludables para prevenir la hipertensión.")
    
    return "\n\n".join(recommendations)

# Función para reemplazar los datos personales y aplicar formato
def replace_personal_data(para, patient):
    # Reemplazar directamente en todo el texto del párrafo (no usar runs)
    para_text = para.text
    para_text = para_text.replace('[NombrePaciente]', f"{patient.nombre} {patient.apellido_p} {patient.apellido_m}")
    para_text = para_text.replace('[FechaNacimiento]', f"{patient.fecha_nacimiento.strftime('%d/%m/%Y')}")
    para_text = para_text.replace('[Sexo]', patient.sexo)
    para_text = para_text.replace('[Edad]', f"{patient.edad} años")
    para_text = para_text.replace('[UUID]', patient.uuid)
    para_text = para_text.replace('[NumeroRegistro]', "191124213115661")
    para_text = para_text.replace('[Reportado]', datetime.now().strftime('%d/%m/%Y - %H:%M UTC'))
    para_text = para_text.replace('[NombreEstudio]', "Análisis Clínico Avanzado")
    para_text = para_text.replace('[Referido]', "Ing. Alejandro López Valdivia")
    para_text = para_text.replace('[NombreInstitucion]', "N/A")

    # Limpiar los runs del párrafo
    para.clear()  # Limpiar el contenido del párrafo

    # Agregar el texto reemplazado con el formato correcto
    run = para.add_run(para_text)
    run.font.size = Pt(10)  # Cambiar el tamaño de la fuente
    run.font.name = 'Aptos (Cuerpo)'  # Cambiar la fuente a una fuente común

# Función para reemplazar los comentarios y aplicar formato
def replace_comments(para, patient):
    para_text = para.text
    para_text = para_text.replace('[ComentarioIMC]', generate_comment_IMC(patient))
    para_text = para_text.replace('[ComentarioPresion]', generate_comment_Presion(patient))
    para_text = para_text.replace('[ComentarioColesterol]', generate_comment_Colesterol(patient))
    para_text = para_text.replace('[ComentarioTrigliceridos]', generate_comment_Trigliceridos(patient))
    para_text = para_text.replace('[ComentarioIndiceAterogenico]', generate_comment_IndiceAterogenico(patient))
    para_text = para_text.replace('[ComentarioGlucosa]', generate_comment_Glucosa(patient))
    para_text = para_text.replace('[ComentarioHbA1c]', generate_comment_HbA1c(patient))
    para_text = para_text.replace('[ComentarioFumador]', generate_comment_Fumador(patient))
    para_text = para_text.replace('[ComentarioAlcoholico]', generate_comment_Alcoholico(patient))
    para_text = para_text.replace('[ComentarioActivo]', generate_comment_Activo(patient))
    para_text = para_text.replace('[ComentarioFamiliaresDiabetes]', generate_comment_FamiliaresDiabetes(patient))
    para_text = para_text.replace('[ComentarioFamiliaresHipertension]', generate_comment_FamiliaresHipertension(patient))
    para_text = para_text.replace('[ComentarioFamiliaresCardiovascular]', generate_comment_FamiliaresCardiovascular(patient))
    para_text = para_text.replace('[ComentarioFamiliaresCerebrovascular]', generate_comment_FamiliaresCerebrovascular(patient))

    # Limpiar los runs del párrafo
    para.clear()  # Limpiar el contenido del párrafo

    # Agregar el texto reemplazado con el formato correcto
    run = para.add_run(para_text)
    run.font.size = Pt(10)  # Cambiar el tamaño de la fuente
    run.font.name = 'Aptos (Cuerpo)'  # Cambiar la fuente a una fuente común

# Función para reemplazar las conclusiones y aplicar formato
def replace_conclusions(para, patient):
    conclusions = generate_conclusions(patient)
    para_text = para.text.replace('[Conclusiones]', conclusions)

    # Limpiar los runs del párrafo
    para.clear()  # Limpiar el contenido del párrafo

    # Agregar el texto reemplazado con el formato correcto
    run = para.add_run(para_text)
    run.font.size = Pt(10)  # Cambiar el tamaño de la fuente
    run.font.name = 'Aptos (Cuerpo)'  # Cambiar la fuente a una fuente común

# Función para reemplazar las recomendaciones y aplicar formato
def replace_recommendations(para, patient):
    recommendations = generate_recommendations(patient)
    para_text = para.text.replace('[Recomendaciones]', recommendations)

    # Limpiar los runs del párrafo
    para.clear()  # Limpiar el contenido del párrafo

    # Agregar el texto reemplazado con el formato correcto
    run = para.add_run(para_text)
    run.font.size = Pt(10)  # Cambiar el tamaño de la fuente
    run.font.name = 'Aptos (Cuerpo)'  # Cambiar la fuente a una fuente común

import os
import subprocess
from docx import Document
from datetime import datetime

# Función para convertir el archivo Word a PDF usando unoconv
def convert_word_to_pdf(input_path, output_path):
    # Usar unoconv para convertir el archivo .docx a .pdf
    command = f"unoconv -f pdf -o {output_path} {input_path}"
    subprocess.run(command, shell=True)

# Función para abrir el PDF automáticamente
def open_pdf(file_path):
    if os.name == 'posix':  # macOS/Linux
        subprocess.run(['open', file_path])  # Para macOS usa 'open', en Linux 'xdg-open'
    elif os.name == 'nt':  # Windows
        os.startfile(file_path)  # Windows

# Función para llenar el reporte con los datos reemplazados y convertir a PDF
def fill_word_template(patient, output_path):
    # Cargar el archivo de plantilla
    doc = Document("./data/datasets/seed.docx")

    # Reemplazar los datos personales en todos los párrafos
    for para in doc.paragraphs:
        replace_personal_data(para, patient)
    
    # Reemplazar los comentarios clínicos
    for para in doc.paragraphs:
        replace_comments(para, patient)

    # Reemplazar las conclusiones
    for para in doc.paragraphs:
        replace_conclusions(para, patient)

    # Reemplazar las recomendaciones
    for para in doc.paragraphs:
        replace_recommendations(para, patient)

    # Guardar el archivo con los datos reemplazados en formato .docx
    doc.save(output_path)

    # Convertir el archivo Word a PDF
    pdf_output_path = output_path.replace('.docx', '.pdf')  # Cambiar extensión a .pdf
    convert_word_to_pdf(output_path, pdf_output_path)  # Realizamos la conversión

    # Abrir el archivo PDF automáticamente
    open_pdf(pdf_output_path)
